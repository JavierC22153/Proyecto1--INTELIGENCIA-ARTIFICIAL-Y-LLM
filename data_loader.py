import os
import pdfplumber
import docx
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_pinecone import PineconeVectorStore
from langchain.docstore.document import Document
from pinecone import Pinecone, ServerlessSpec

load_dotenv()

pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))

if os.getenv("INDEX_NAME") not in pc.list_indexes().names():
    pc.create_index(
        name=os.getenv("INDEX_NAME"),
        dimension=1536,
        metric='cosine',
        spec=ServerlessSpec(
            cloud='aws',
            region=os.getenv("PINECONE_ENVIRONMENT")
        )
    )

embeddings = OpenAIEmbeddings(model="text-embedding-3-small")

def load_txt(file) -> str:
    return file.read().decode("utf-8")

def load_pdf(file) -> str:
    content = ""
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            content += page.extract_text()
    return content

def load_docx(file) -> str:
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def ingest_docs(file) -> None:
    if file.name.endswith(".txt"):
        content = load_txt(file)
    elif file.name.endswith(".pdf"):
        content = load_pdf(file)
    elif file.name.endswith(".docx"):
        content = load_docx(file)
    else:
        raise ValueError("Formato de archivo no soportado. Usa .txt, .pdf o .docx")

    doc = Document(page_content=content, metadata={"source": file.name})

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=200)
    docs = text_splitter.split_documents([doc])

    print(f"Añadiendo {len(docs)} documentos a Pinecone.")

    PineconeVectorStore.from_documents(
        docs, embeddings, index_name=os.getenv("INDEX_NAME")
    )

    print("Documentos locales cargados en Pinecone.")